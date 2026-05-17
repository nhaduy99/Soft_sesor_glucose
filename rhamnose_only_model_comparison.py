import base64
import math
import os
import warnings
from pathlib import Path

import numpy as np
import pandas as pd

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch

from docx import Document
from docx.shared import Inches
from sklearn.base import clone
from sklearn.cross_decomposition import PLSRegression
from sklearn.decomposition import PCA
from sklearn.dummy import DummyRegressor
from sklearn.ensemble import ExtraTreesRegressor, RandomForestRegressor
from sklearn.feature_selection import SelectKBest, f_regression
from sklearn.impute import SimpleImputer
from sklearn.kernel_ridge import KernelRidge
from sklearn.linear_model import ElasticNet, Lasso, LinearRegression, Ridge
from sklearn.metrics import mean_absolute_error, mean_squared_error, r2_score
from sklearn.model_selection import GroupShuffleSplit
from sklearn.neighbors import KNeighborsRegressor
from sklearn.neural_network import MLPRegressor
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import FunctionTransformer, StandardScaler
from sklearn.svm import SVR

from train_monosaccharide_softsensor import (
    FULL_CSV,
    INTERPRETABLE_CSV,
    MERGE_KEY_COLUMNS,
    feature_columns,
    merge_modalities,
    parse_targets,
    read_csv,
    safe_float,
    target_rows,
    write_csv,
)

try:
    from xgboost import XGBRegressor
except Exception:
    XGBRegressor = None


ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "rhamnose_model_comparison"
FIG_DIR = OUT_DIR / "figures"
RAMAN_PREPROCESSED_CSV = ROOT / "features" / "raman_preprocessed_features.csv"
PARAFAC_CSV = ROOT / "features" / "eem_parafac_scores_exclude_rha5.csv"
TARGET = "rhamnose_gL"
EXCLUDE_RHA5 = True
RANDOM_STATE = 17

warnings.filterwarnings("ignore", category=UserWarning)
warnings.filterwarnings("ignore", category=FutureWarning)
warnings.filterwarnings("ignore", category=RuntimeWarning)


NATURE = {
    "blue": "#0072B2",
    "green": "#009E73",
    "orange": "#D55E00",
    "purple": "#CC79A7",
    "sky": "#56B4E9",
    "black": "#111111",
    "grey": "#6B7280",
    "grid": "#E5E7EB",
}


def key_for(row):
    return tuple(row.get(col, "") for col in MERGE_KEY_COLUMNS)


def is_rha5(row):
    label = str(row.get("legend_treatment_label", "")).strip().lower().replace(" ", "")
    return label == "rha(5)" and safe_float(row.get(TARGET)) == 5.0


def add_targets_if_needed(rows):
    out = []
    for row in rows:
        new = dict(row)
        if TARGET not in new or str(new.get(TARGET, "")).strip() == "":
            targets, source = parse_targets(new.get("legend_treatment_label", ""))
            new["target_source"] = source
            for target, value in targets.items():
                new[target] = f"{value:.8g}"
        out.append(new)
    return out


def load_rows(path, parse_target_labels=True):
    rows = read_csv(path)
    if parse_target_labels:
        rows = target_rows(rows)
    else:
        rows = add_targets_if_needed(rows)
    rows = merge_modalities(rows)
    if EXCLUDE_RHA5:
        rows = [row for row in rows if not is_rha5(row)]
    return rows


def merge_right(left_rows, right_rows, prefixes, passthrough=()):
    by_key = {key_for(row): row for row in right_rows}
    merged = []
    for left in left_rows:
        right = by_key.get(key_for(left))
        if not right:
            continue
        row = dict(left)
        for key, value in right.items():
            if any(key.startswith(prefix) for prefix in prefixes) or key in passthrough:
                row[key] = value
        merged.append(row)
    return merged


def feature_cols(rows, prefixes):
    fields = list(rows[0].keys()) if rows else []
    return [col for col in fields if any(col.startswith(prefix) for prefix in prefixes)]


def matrix_from_rows(rows, cols):
    x_rows, y_vals, kept = [], [], []
    for row in rows:
        y = safe_float(row.get(TARGET))
        if not math.isfinite(y):
            continue
        vals = [safe_float(row.get(col)) for col in cols]
        if not any(math.isfinite(v) for v in vals):
            continue
        x_rows.append(vals)
        y_vals.append(y)
        kept.append(row)
    return np.asarray(x_rows, dtype=float), np.asarray(y_vals, dtype=float), kept


def groups_for(rows):
    groups = []
    for row in rows:
        parts = [
            row.get("metadata_experiment", ""),
            row.get("metadata_plate", ""),
            row.get("metadata_well", ""),
            row.get("legend_treatment_label", ""),
        ]
        groups.append("|".join(str(part) for part in parts))
    return np.asarray(groups)


def finite_k(k, x):
    return max(1, min(k, x.shape[1]))


def select_step(k):
    return SelectKBest(score_func=f_regression, k=k)


def pca_regressor(n_components=5):
    return Pipeline(
        [
            ("pca", PCA(n_components=n_components, random_state=RANDOM_STATE)),
            ("ridge", Ridge(alpha=1.0)),
        ]
    )


def model_specs(n_features):
    max_k = finite_k(192, np.zeros((1, n_features)))
    specs = [
        ("Statistics", "Mean baseline", "mean", DummyRegressor(strategy="mean"), min(12, max_k)),
        ("Statistics", "Linear regression", "ols", LinearRegression(), min(96, max_k)),
        ("Statistics", "Ridge", "alpha=10", Ridge(alpha=10.0), max_k),
        ("Statistics", "Lasso", "alpha=0.001", Lasso(alpha=0.001, max_iter=10000, random_state=RANDOM_STATE), max_k),
        (
            "Statistics",
            "Elastic Net",
            "alpha=0.001,l1_ratio=0.2",
            ElasticNet(alpha=0.001, l1_ratio=0.2, max_iter=10000, random_state=RANDOM_STATE),
            max_k,
        ),
        ("Chemometrics", "PLSR", "n_components=2", PLSRegression(n_components=2), max_k),
        ("Chemometrics", "PLSR", "n_components=5", PLSRegression(n_components=5), max_k),
        ("Chemometrics", "PCR + Ridge", "n_components=5", pca_regressor(5), max_k),
        ("ML", "SVR linear", "C=1,epsilon=0.05", SVR(kernel="linear", C=1.0, epsilon=0.05), max_k),
        ("ML", "SVR RBF", "C=10,epsilon=0.05", SVR(kernel="rbf", C=10.0, epsilon=0.05), max_k),
        ("ML", "Weighted kNN", "k=3,distance", KNeighborsRegressor(n_neighbors=3, weights="distance", metric="manhattan"), min(96, max_k)),
        ("ML", "Kernel ridge", "rbf,alpha=0.1", KernelRidge(kernel="rbf", alpha=0.1), max_k),
        (
            "ML",
            "Random forest",
            "n=250,max_depth=4",
            RandomForestRegressor(n_estimators=250, max_depth=4, random_state=RANDOM_STATE, n_jobs=-1),
            min(96, max_k),
        ),
        (
            "ML",
            "Extra trees",
            "n=250,max_depth=4",
            ExtraTreesRegressor(n_estimators=250, max_depth=4, random_state=RANDOM_STATE, n_jobs=-1),
            min(96, max_k),
        ),
        (
            "DL",
            "ANN MLP",
            "hidden=(32,16),early_stop",
            MLPRegressor(
                hidden_layer_sizes=(32, 16),
                activation="relu",
                alpha=0.001,
                learning_rate_init=0.001,
                max_iter=500,
                early_stopping=True,
                random_state=RANDOM_STATE,
            ),
            min(96, max_k),
        ),
    ]
    if XGBRegressor is not None:
        specs.append(
            (
                "ML",
                "XGBoost",
                "n=250,depth=2,eta=0.05",
                XGBRegressor(
                    n_estimators=250,
                    max_depth=2,
                    learning_rate=0.05,
                    subsample=0.9,
                    colsample_bytree=0.9,
                    objective="reg:squarederror",
                    random_state=RANDOM_STATE,
                    n_jobs=1,
                ),
                min(192, max_k),
            )
        )
    return specs


def make_pipeline(estimator, k, n_features):
    k = finite_k(k, np.zeros((1, n_features)))
    return Pipeline(
        [
            ("impute", SimpleImputer(strategy="median")),
            ("nan_guard", FunctionTransformer(lambda x: np.nan_to_num(x, nan=0.0, posinf=0.0, neginf=0.0))),
            ("select", select_step(k)),
            ("scale", StandardScaler()),
            ("model", clone(estimator)),
        ]
    )


def eval_feature_set(name, rows, cols):
    x, y, kept = matrix_from_rows(rows, cols)
    if len(kept) < 12:
        return [], []
    groups = groups_for(kept)
    splitter = GroupShuffleSplit(n_splits=5, test_size=0.25, random_state=RANDOM_STATE)
    metric_rows, pred_rows = [], []
    for split_id, (train_idx, test_idx) in enumerate(splitter.split(x, y, groups)):
        x_train, x_test = x[train_idx], x[test_idx]
        y_train, y_test = y[train_idx], y[test_idx]
        for category, model_name, config, estimator, k in model_specs(x.shape[1]):
            pipe = make_pipeline(estimator, k, x.shape[1])
            try:
                pipe.fit(x_train, y_train)
                pred = np.asarray(pipe.predict(x_test), dtype=float).reshape(-1)
            except Exception as exc:
                metric_rows.append(
                    {
                        "feature_strategy": name,
                        "model_category": category,
                        "model": model_name,
                        "config": config,
                        "split": split_id,
                        "rmse": "",
                        "mae": "",
                        "r2": "",
                        "status": f"failed: {type(exc).__name__}",
                    }
                )
                continue
            rmse = float(np.sqrt(mean_squared_error(y_test, pred)))
            mae = float(mean_absolute_error(y_test, pred))
            r2 = float(r2_score(y_test, pred)) if len(np.unique(y_test)) > 1 else math.nan
            metric_rows.append(
                {
                    "feature_strategy": name,
                    "model_category": category,
                    "model": model_name,
                    "config": config,
                    "split": split_id,
                    "n_samples": len(kept),
                    "n_train": len(train_idx),
                    "n_test": len(test_idx),
                    "n_features_raw": len(cols),
                    "n_features_selected": finite_k(k, x),
                    "rmse": rmse,
                    "mae": mae,
                    "r2": r2,
                    "status": "ok",
                }
            )
            if split_id == 0:
                for local_idx, row_idx in enumerate(test_idx):
                    src = kept[row_idx]
                    pred_rows.append(
                        {
                            "feature_strategy": name,
                            "model_category": category,
                            "model": model_name,
                            "config": config,
                            "sample_id": src.get("sample_id", ""),
                            "batch": src.get("batch", ""),
                            "metadata_plate": src.get("metadata_plate", ""),
                            "metadata_well": src.get("metadata_well", ""),
                            "legend_treatment_label": src.get("legend_treatment_label", ""),
                            "y_true": float(y_test[local_idx]),
                            "y_pred": float(pred[local_idx]),
                            "residual": float(y_test[local_idx] - pred[local_idx]),
                        }
                    )
    return metric_rows, pred_rows


def aggregate_metrics(metrics):
    df = pd.DataFrame(metrics)
    ok = df[df["status"].eq("ok")].copy()
    for col in ("rmse", "mae", "r2"):
        ok[col] = pd.to_numeric(ok[col], errors="coerce")
    summary = (
        ok.groupby(["feature_strategy", "model_category", "model", "config"], dropna=False)
        .agg(
            n_splits=("split", "count"),
            n_samples=("n_samples", "max"),
            n_features_raw=("n_features_raw", "max"),
            n_features_selected=("n_features_selected", "max"),
            mean_rmse=("rmse", "mean"),
            std_rmse=("rmse", "std"),
            mean_mae=("mae", "mean"),
            mean_r2=("r2", "mean"),
        )
        .reset_index()
        .sort_values(["mean_rmse", "mean_mae"], ascending=True)
    )
    return summary


def save_plot(path):
    plt.tight_layout()
    plt.savefig(path, dpi=220, bbox_inches="tight")
    plt.close()


def style_axes(ax):
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.grid(True, color=NATURE["grid"], linewidth=0.8, zorder=0)
    ax.tick_params(colors=NATURE["black"], labelsize=8)
    ax.xaxis.label.set_size(9)
    ax.yaxis.label.set_size(9)
    ax.title.set_size(10)


def make_figures(summary, pred_df):
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    plt.figure(figsize=(10, 4.8))
    ax = plt.gca()
    ax.set_axis_off()
    boxes = [
        (0.04, 0.66, "Raw spectroscopy", "EEM matrices\nRaman spectra"),
        (0.29, 0.78, "EEM path", "Unfolded cells\nInterpretable hotspots\nPARAFAC scores"),
        (0.29, 0.52, "Raman path", "Raw windows\nFull spectrum\nProcessed Raman"),
        (0.55, 0.66, "Feature strategies", "EEM-only\nRaman-only\nMid/full fusion"),
        (0.76, 0.76, "Model families", "Statistics\nPLS/PCR\nML/XGBoost\nANN MLP"),
        (0.76, 0.48, "Validation", "Grouped splits\nRMSE, MAE, R2\nPred-vs-true"),
    ]
    for x, y, title, body in boxes:
        patch = FancyBboxPatch(
            (x, y - 0.11),
            0.18,
            0.18,
            boxstyle="round,pad=0.018,rounding_size=0.012",
            linewidth=1,
            edgecolor="#C7CDD4",
            facecolor="#FFFFFF",
        )
        ax.add_patch(patch)
        ax.text(x + 0.09, y + 0.025, title, ha="center", va="center", fontsize=10, fontweight="bold", color=NATURE["black"])
        ax.text(x + 0.09, y - 0.055, body, ha="center", va="center", fontsize=8.5, color="#374151")
    arrows = [
        ((0.22, 0.68), (0.29, 0.79)),
        ((0.22, 0.64), (0.29, 0.53)),
        ((0.47, 0.78), (0.55, 0.69)),
        ((0.47, 0.52), (0.55, 0.65)),
        ((0.73, 0.68), (0.76, 0.77)),
        ((0.73, 0.65), (0.76, 0.49)),
    ]
    for start, end in arrows:
        ax.add_patch(FancyArrowPatch(start, end, arrowstyle="-|>", mutation_scale=12, color=NATURE["grey"], linewidth=1.2))
    ax.text(0.5, 0.17, "Scientific interpretation: Raman provides direct carbohydrate-sensitive signal; EEM/PARAFAC provides matrix-state and fluorescence-component context.", ha="center", fontsize=9, color="#374151")
    ax.set_xlim(0, 1)
    ax.set_ylim(0.2, 1.0)
    plt.title("Rhamnose-only modelling pipeline", fontsize=12, fontweight="bold")
    save_plot(FIG_DIR / "pipeline_diagram.png")

    top = summary.head(20).copy()
    plt.figure(figsize=(7.2, 5.0))
    labels = [f"{r.feature_strategy}\n{r.model}" for r in top.itertuples()]
    plt.barh(range(len(top)), top["mean_rmse"], color=NATURE["blue"], alpha=0.9)
    plt.yticks(range(len(top)), labels)
    plt.gca().invert_yaxis()
    plt.xlabel("Cross-validated RMSE (g/L)")
    plt.title("Top rhamnose model/feature strategies")
    style_axes(plt.gca())
    save_plot(FIG_DIR / "top20_rmse.png")

    best_by_strategy = summary.loc[summary.groupby("feature_strategy")["mean_rmse"].idxmin()].sort_values("mean_rmse")
    plt.figure(figsize=(7.2, 4.2))
    plt.bar(best_by_strategy["feature_strategy"], best_by_strategy["mean_rmse"], color=NATURE["green"], alpha=0.9)
    plt.ylabel("Best RMSE within strategy (g/L)")
    plt.xlabel("Input strategy")
    plt.xticks(rotation=35, ha="right")
    plt.title("Best rhamnose performance by input strategy")
    style_axes(plt.gca())
    save_plot(FIG_DIR / "strategy_rmse.png")

    best_by_category = summary.loc[summary.groupby("model_category")["mean_rmse"].idxmin()].sort_values("mean_rmse")
    plt.figure(figsize=(5.5, 3.8))
    plt.bar(best_by_category["model_category"], best_by_category["mean_rmse"], color=NATURE["orange"], alpha=0.9)
    plt.ylabel("Best RMSE within model category (g/L)")
    plt.xlabel("Model family")
    plt.title("Best rhamnose performance by model family")
    style_axes(plt.gca())
    save_plot(FIG_DIR / "category_rmse.png")

    best = summary.iloc[0]
    rows = pred_df[
        pred_df["feature_strategy"].eq(best["feature_strategy"])
        & pred_df["model"].eq(best["model"])
        & pred_df["config"].eq(best["config"])
    ].copy()
    plt.figure(figsize=(4.4, 4.2))
    lo = min(rows["y_true"].min(), rows["y_pred"].min(), 0)
    hi = max(rows["y_true"].max(), rows["y_pred"].max())
    pad = 0.08 * (hi - lo if hi > lo else 1)
    plt.plot([lo - pad, hi + pad], [lo - pad, hi + pad], color=NATURE["grey"], linestyle="--", linewidth=1)
    plt.scatter(rows["y_true"], rows["y_pred"], s=34, color=NATURE["blue"], alpha=0.75, edgecolor="white", linewidth=0.4)
    plt.xlabel("True rhamnose concentration (g/L)")
    plt.ylabel("Predicted rhamnose concentration (g/L)")
    plt.title("Best model: predicted vs true")
    style_axes(plt.gca())
    save_plot(FIG_DIR / "best_pred_vs_true.png")

    plt.figure(figsize=(4.8, 3.6))
    plt.axhline(0, color=NATURE["grey"], linestyle="--", linewidth=1)
    plt.scatter(rows["y_true"], rows["residual"], s=34, color=NATURE["purple"], alpha=0.75, edgecolor="white", linewidth=0.4)
    plt.xlabel("True rhamnose concentration (g/L)")
    plt.ylabel("Residual: true - predicted (g/L)")
    plt.title("Best model residual diagnostic")
    style_axes(plt.gca())
    save_plot(FIG_DIR / "best_residuals.png")

    return best, best_by_strategy, best_by_category


def img_tag(path, caption):
    data = base64.b64encode(path.read_bytes()).decode("ascii")
    return f'<figure><img src="data:image/png;base64,{data}" alt="{caption}"><figcaption>{caption}</figcaption></figure>'


def write_html(summary, best_by_strategy, best_by_category, best):
    top_rows = "\n".join(
        f"<tr><td>{r.feature_strategy}</td><td>{r.model_category}</td><td>{r.model}</td><td>{r.config}</td>"
        f"<td>{r.mean_rmse:.4f}</td><td>{r.std_rmse:.4f}</td><td>{r.mean_mae:.4f}</td><td>{r.mean_r2:.3f}</td></tr>"
        for r in summary.head(30).itertuples()
    )
    strategy_rows = "\n".join(
        f"<tr><td>{r.feature_strategy}</td><td>{r.model_category}</td><td>{r.model}</td><td>{r.mean_rmse:.4f}</td><td>{r.mean_r2:.3f}</td></tr>"
        for r in best_by_strategy.itertuples()
    )
    category_rows = "\n".join(
        f"<tr><td>{r.model_category}</td><td>{r.feature_strategy}</td><td>{r.model}</td><td>{r.mean_rmse:.4f}</td><td>{r.mean_r2:.3f}</td></tr>"
        for r in best_by_category.itertuples()
    )
    html = f"""<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8">
<title>Rhamnose-Only Soft-Sensor Model Comparison</title>
<style>
body {{ margin:0; background:#f7f7f7; color:#111; font-family:Arial, Helvetica, sans-serif; line-height:1.55; }}
main {{ max-width:1180px; margin:0 auto; padding:32px 24px 64px; }}
h1 {{ font-size:30px; margin-bottom:4px; }}
h2 {{ margin-top:34px; font-size:21px; }}
.panel {{ background:#fff; border:1px solid #ddd; border-radius:4px; padding:18px; margin:16px 0; }}
.grid {{ display:grid; grid-template-columns:repeat(auto-fit, minmax(420px, 1fr)); gap:16px; }}
table {{ width:100%; border-collapse:collapse; font-size:13px; }}
th, td {{ border-bottom:1px solid #e5e7eb; padding:7px 8px; text-align:left; vertical-align:top; }}
th {{ background:#eef3f8; font-weight:700; }}
figure {{ margin:0; background:#fff; border:1px solid #ddd; border-radius:4px; padding:12px; }}
img {{ max-width:100%; display:block; }}
figcaption {{ font-size:12px; color:#4b5563; margin-top:8px; }}
.note {{ color:#4b5563; }}
code {{ background:#eef3f8; padding:1px 4px; border-radius:3px; }}
</style>
</head>
<body><main>
<h1>Rhamnose-Only Soft-Sensor Model Comparison</h1>
<p class="note">Academic comparison of input strategies and model families for rhamnose concentration prediction. All <code>Rha (5)</code> examples are excluded to match the current corrected sensitivity-analysis workflow.</p>

<section class="panel">
<h2>Executive Result</h2>
<p>The best cross-validated result in this rhamnose-only comparison is <strong>{best['feature_strategy']} + {best['model']}</strong> ({best['config']}), with mean RMSE <strong>{best['mean_rmse']:.4f} g/L</strong> and mean R2 <strong>{best['mean_r2']:.3f}</strong>.</p>
<p>Interpretation: rhamnose is not strongly fluorescent, so Raman and fused representations are expected to carry the most direct chemistry. EEM-derived features are still useful as process-state or matrix-context signals, especially when fused with Raman.</p>
</section>

<h2>Pipeline</h2>
<section class="panel">
<p><strong>Raw inputs:</strong> EEM matrices, Raman spectra, processed Raman features, EEM PARAFAC scores, and fused feature tables.</p>
<p><strong>Preprocessing:</strong> missing-value imputation, variance filtering, univariate feature selection inside each training fold, standard scaling, and grouped train/test splits by plate/well/treatment label. This prevents exact replicate groups from being split casually across train and test.</p>
<p><strong>Models:</strong> statistical baselines and linear models, chemometric PLS/PCR models, classical ML models, XGBoost, and a neural-network MLP regressor. All models use identical grouped splits for fair comparison.</p>
<p><strong>Outputs:</strong> RMSE, MAE, R2, best-strategy tables, predicted-vs-true plots, residual diagnostics, and academic interpretation.</p>
</section>

<h2>Figures</h2>
<div class="grid">
{img_tag(FIG_DIR / 'top20_rmse.png', 'Figure 1. Top 20 rhamnose model and input-strategy combinations ranked by grouped cross-validated RMSE. Lower RMSE is better.')}
{img_tag(FIG_DIR / 'pipeline_diagram.png', 'Figure 2. End-to-end rhamnose-only modelling pipeline showing raw spectroscopy inputs, feature strategies, model families, and grouped validation outputs.')}
{img_tag(FIG_DIR / 'strategy_rmse.png', 'Figure 3. Best model within each input strategy. This directly compares EEM-only, Raman-only, processed, featured, PARAFAC, and fusion inputs.')}
{img_tag(FIG_DIR / 'category_rmse.png', 'Figure 4. Best model within each model family: statistical, chemometric, machine learning, and neural-network style models.')}
{img_tag(FIG_DIR / 'best_pred_vs_true.png', 'Figure 5. Predicted versus true rhamnose concentration for the best model on held-out grouped split 0. Points close to the diagonal indicate better calibration.')}
{img_tag(FIG_DIR / 'best_residuals.png', 'Figure 6. Residual diagnostic for the best model. Random scatter around zero is preferred; systematic trends imply concentration-dependent bias.')}
</div>

<h2>Best Models by Input Strategy</h2>
<section class="panel"><table><tr><th>Input strategy</th><th>Model family</th><th>Best model</th><th>RMSE</th><th>R2</th></tr>{strategy_rows}</table></section>

<h2>Best Models by Model Family</h2>
<section class="panel"><table><tr><th>Model family</th><th>Input strategy</th><th>Best model</th><th>RMSE</th><th>R2</th></tr>{category_rows}</table></section>

<h2>Top Ranked Configurations</h2>
<section class="panel"><table><tr><th>Input strategy</th><th>Model family</th><th>Model</th><th>Config</th><th>Mean RMSE</th><th>SD RMSE</th><th>Mean MAE</th><th>Mean R2</th></tr>{top_rows}</table></section>

<h2>Analysis for Academic Audience</h2>
<section class="panel">
<p><strong>Strengths.</strong> This comparison tests multiple chemically distinct input representations rather than assuming one modality is sufficient. Raman features target direct carbohydrate vibrational information; EEM features provide indirect matrix/process-state information; PARAFAC scores provide interpretable latent fluorescence components; fusion models test whether direct and indirect signals are complementary.</p>
<p><strong>Weaknesses.</strong> Current labels are still standards/known spike concentrations parsed from treatment names, not final quantitative culture-sample HPLC concentrations. The sample size is modest, so flexible models such as XGBoost, random forests, and MLPs can overfit. Results should therefore be read as calibration/sensitivity evidence, not as final biological deployment performance.</p>
<p><strong>Recommended interpretation.</strong> Prefer the simplest model within one standard error of the best RMSE for publication claims. If a fusion model wins, the result supports the scientific hypothesis that Raman supplies direct rhamnose chemistry while EEM supplies contextual fluorescence/process-state information.</p>
</section>
</main></body></html>"""
    (OUT_DIR / "rhamnose_model_comparison_report.html").write_text(html, encoding="utf-8")


def write_docx(summary, best_by_strategy, best_by_category, best):
    doc = Document()
    doc.add_heading("Rhamnose-Only Soft-Sensor Model Comparison", 0)
    doc.add_paragraph(
        "Academic comparison of input strategies and model families for rhamnose concentration prediction. "
        "All Rha (5) examples are excluded to match the current corrected sensitivity-analysis workflow."
    )
    doc.add_heading("Executive Result", level=1)
    doc.add_paragraph(
        f"The best cross-validated result is {best['feature_strategy']} + {best['model']} ({best['config']}), "
        f"with mean RMSE {best['mean_rmse']:.4f} g/L and mean R2 {best['mean_r2']:.3f}."
    )
    doc.add_heading("Pipeline", level=1)
    for text in [
        "Raw inputs: EEM matrices, Raman spectra, processed Raman features, EEM PARAFAC scores, and fused feature tables.",
        "Preprocessing: missing-value imputation, variance filtering, univariate feature selection inside each training fold, standard scaling, and grouped train/test splits by plate/well/treatment label.",
        "Models: statistical baselines and linear models, chemometric PLS/PCR models, classical ML models, XGBoost, and a neural-network MLP regressor.",
        "Outputs: RMSE, MAE, R2, best-strategy tables, predicted-vs-true plots, residual diagnostics, and academic interpretation.",
    ]:
        doc.add_paragraph(text, style="List Bullet")
    doc.add_heading("Figures", level=1)
    for filename, caption in [
        ("top20_rmse.png", "Figure 1. Top 20 rhamnose model and input-strategy combinations ranked by grouped cross-validated RMSE."),
        ("pipeline_diagram.png", "Figure 2. End-to-end rhamnose-only modelling pipeline."),
        ("strategy_rmse.png", "Figure 3. Best model within each input strategy."),
        ("category_rmse.png", "Figure 4. Best model within each model family."),
        ("best_pred_vs_true.png", "Figure 5. Predicted versus true rhamnose concentration for the best model."),
        ("best_residuals.png", "Figure 6. Residual diagnostic for the best model."),
    ]:
        doc.add_picture(str(FIG_DIR / filename), width=Inches(6.2))
        doc.add_paragraph(caption)
    doc.add_heading("Best Models by Input Strategy", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for idx, text in enumerate(["Input strategy", "Model family", "Best model", "RMSE", "R2"]):
        table.rows[0].cells[idx].text = text
    for row in best_by_strategy.itertuples():
        cells = table.add_row().cells
        cells[0].text = str(row.feature_strategy)
        cells[1].text = str(row.model_category)
        cells[2].text = str(row.model)
        cells[3].text = f"{row.mean_rmse:.4f}"
        cells[4].text = f"{row.mean_r2:.3f}"
    doc.add_heading("Top Ranked Configurations", level=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    for idx, text in enumerate(["Input strategy", "Family", "Model", "Config", "RMSE", "R2"]):
        table.rows[0].cells[idx].text = text
    for row in summary.head(20).itertuples():
        cells = table.add_row().cells
        cells[0].text = str(row.feature_strategy)
        cells[1].text = str(row.model_category)
        cells[2].text = str(row.model)
        cells[3].text = str(row.config)
        cells[4].text = f"{row.mean_rmse:.4f}"
        cells[5].text = f"{row.mean_r2:.3f}"
    doc.add_heading("Academic Analysis", level=1)
    doc.add_paragraph(
        "The comparison is strongest as a calibration and method-selection study. Raman should be interpreted as the most direct rhamnose-sensitive modality, while EEM and PARAFAC features provide indirect matrix-state and fluorescence-component information."
    )
    doc.add_paragraph(
        "The main limitation is target quality: the current labels are standards and known spike concentrations parsed from treatment labels. Quantitative culture-sample HPLC concentrations are still required before claiming biological deployment performance."
    )
    doc.save(OUT_DIR / "rhamnose_model_comparison_report.docx")


def write_readme(best):
    text = f"""# Rhamnose-Only Model Comparison

This folder compares rhamnose prediction strategies using EEM, Raman, processed Raman, EEM PARAFAC, and fused inputs.

Current best result:

- Input strategy: `{best['feature_strategy']}`
- Model: `{best['model']}`
- Config: `{best['config']}`
- Mean grouped CV RMSE: `{best['mean_rmse']:.4f}` g/L
- Mean grouped CV R2: `{best['mean_r2']:.3f}`

Main files:

- `rhamnose_model_metrics_by_split.csv`
- `rhamnose_model_metrics_summary.csv`
- `rhamnose_model_predictions_split0.csv`
- `rhamnose_model_comparison_report.html`
- `rhamnose_model_comparison_report.docx`
- `figures/`

Run command:

```powershell
conda run -n base python rhamnose_only_model_comparison.py
```
"""
    (OUT_DIR / "README.md").write_text(text, encoding="utf-8")


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)

    interp = load_rows(INTERPRETABLE_CSV, parse_target_labels=True)
    full = load_rows(FULL_CSV, parse_target_labels=True)
    raman_processed = load_rows(RAMAN_PREPROCESSED_CSV, parse_target_labels=False)
    parafac = load_rows(PARAFAC_CSV, parse_target_labels=False)

    rp_parafac = merge_right(raman_processed, parafac, ("parafac_score",), ("parafac_selected_rank",))
    interp_parafac = merge_right(interp, parafac, ("parafac_score",), ("parafac_selected_rank",))
    full_rp = merge_right(full, raman_processed, ("rp_",), ("preprocessing_config",))
    full_parafac = merge_right(full, parafac, ("parafac_score",), ("parafac_selected_rank",))

    strategies = [
        ("EEM interpretable", interp, feature_columns(list(interp[0].keys()), "eem_interpretable")),
        ("Raman interpretable", interp, feature_columns(list(interp[0].keys()), "raman_interpretable")),
        ("EEM + Raman interpretable", interp, feature_columns(list(interp[0].keys()), "fusion_interpretable")),
        ("EEM unfolded full", full, feature_columns(list(full[0].keys()), "eem_full")),
        ("Raman full", full, feature_columns(list(full[0].keys()), "raman_full")),
        ("EEM + Raman full", full, feature_columns(list(full[0].keys()), "fusion_full")),
        ("Processed Raman", raman_processed, feature_cols(raman_processed, ("rp_",))),
        ("EEM PARAFAC scores", parafac, feature_cols(parafac, ("parafac_score",))),
        ("Processed Raman + PARAFAC", rp_parafac, feature_cols(rp_parafac, ("rp_", "parafac_score"))),
        ("EEM interpretable + PARAFAC", interp_parafac, feature_columns(list(interp_parafac[0].keys()), "eem_interpretable") + feature_cols(interp_parafac, ("parafac_score",))),
        ("EEM full + processed Raman", full_rp, feature_columns(list(full_rp[0].keys()), "eem_full") + feature_cols(full_rp, ("rp_",))),
        ("EEM full + PARAFAC", full_parafac, feature_columns(list(full_parafac[0].keys()), "eem_full") + feature_cols(full_parafac, ("parafac_score",))),
    ]

    all_metrics, all_predictions = [], []
    for strategy, rows, cols in strategies:
        if not rows or not cols:
            continue
        metrics, predictions = eval_feature_set(strategy, rows, cols)
        all_metrics.extend(metrics)
        all_predictions.extend(predictions)
        print(f"{strategy}: {len(metrics)} split-model rows")

    metrics_df = pd.DataFrame(all_metrics)
    pred_df = pd.DataFrame(all_predictions)
    summary = aggregate_metrics(all_metrics)
    metrics_df.to_csv(OUT_DIR / "rhamnose_model_metrics_by_split.csv", index=False)
    summary.to_csv(OUT_DIR / "rhamnose_model_metrics_summary.csv", index=False)
    pred_df.to_csv(OUT_DIR / "rhamnose_model_predictions_split0.csv", index=False)

    best, best_by_strategy, best_by_category = make_figures(summary, pred_df)
    write_html(summary, best_by_strategy, best_by_category, best)
    write_docx(summary, best_by_strategy, best_by_category, best)
    write_readme(best)
    print(f"Best: {best['feature_strategy']} + {best['model']} RMSE={best['mean_rmse']:.4f}")
    print(f"Wrote {OUT_DIR}")


if __name__ == "__main__":
    main()
